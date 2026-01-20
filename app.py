import hashlib
import io
import json
import os
import sys
from typing import Any, Iterable, Tuple

import langextract as lx
import streamlit as st
from dotenv import load_dotenv
from langextract.core import data as lx_data


load_dotenv()


def _build_extraction(
    extraction_class: str,
    value: Any,
) -> lx_data.Extraction:
    attributes = None
    text_value = value
    if isinstance(value, dict):
        text_value = value.get("text") or value.get("value") or value.get(
            "extraction_text"
        )
        attributes = value.get("attributes")
        if text_value is None:
            text_value = json.dumps(value, ensure_ascii=True)
    if text_value is None:
        text_value = ""
    return lx_data.Extraction(
        extraction_class=str(extraction_class),
        extraction_text=str(text_value),
        attributes=attributes,
    )


def _coerce_extractions(output: Any) -> list[lx_data.Extraction]:
    if output is None:
        return []
    if isinstance(output, dict) and "extractions" in output:
        return _coerce_extractions(output.get("extractions"))
    if isinstance(output, list):
        extractions: list[lx_data.Extraction] = []
        for item in output:
            if not isinstance(item, dict):
                raise ValueError(
                    "Output list must contain extraction objects."
                )
            extraction_class = (
                item.get("extraction_class")
                or item.get("class")
                or item.get("label")
            )
            extraction_text = (
                item.get("extraction_text")
                or item.get("text")
                or item.get("value")
            )
            if extraction_class is None:
                raise ValueError(
                    "Extraction object missing extraction_class."
                )
            extractions.append(
                _build_extraction(extraction_class, extraction_text)
            )
        return extractions
    if isinstance(output, dict):
        extractions = []
        for key, value in output.items():
            if isinstance(value, list):
                for entry in value:
                    extractions.append(_build_extraction(key, entry))
            else:
                extractions.append(_build_extraction(key, value))
        return extractions
    raise ValueError("Output must be a dict or list.")


def _deserialize_examples(raw: str) -> list[lx_data.ExampleData]:
    if not raw.strip():
        return []
    parsed = json.loads(raw)
    if not isinstance(parsed, list):
        raise ValueError("Examples must be a JSON array.")
    examples: list[lx_data.ExampleData] = []
    for item in parsed:
        if not isinstance(item, dict):
            raise ValueError("Each example must be a JSON object.")
        text = item.get("input") or item.get("text")
        if text is None:
            raise ValueError("Each example must include input/text.")
        if "extractions" in item:
            extractions = _coerce_extractions(item.get("extractions"))
        else:
            extractions = _coerce_extractions(item.get("output"))
        examples.append(lx_data.ExampleData(text=text, extractions=extractions))
    return examples


def _result_to_jsonable(result: Any) -> Any:
    if hasattr(result, "to_dict"):
        return result.to_dict()
    if hasattr(result, "model_dump"):
        return result.model_dump()
    if hasattr(result, "__dict__"):
        return result.__dict__
    return result


def _format_table(rows: Iterable[Iterable[str]]) -> str:
    lines = []
    for row in rows:
        cells = [(cell or "").replace("\n", " ").strip() for cell in row]
        lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_pdf_text(data: bytes) -> str:
    try:
        import pdfplumber
    except Exception:  # noqa: BLE001
        pdfplumber = None

    if pdfplumber is not None:
        parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                parts.append(f"=== Page {page_index} ===")
                try:
                    text = page.extract_text(layout=True)
                except TypeError:
                    text = page.extract_text()
                if text:
                    parts.append(text)
                tables = page.extract_tables() or []
                for table_index, table in enumerate(tables, start=1):
                    parts.append(f"--- Table {page_index}.{table_index} ---")
                    parts.append(_format_table(table))
        return "\n".join(p for p in parts if p)

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _ocr_pdf_text(data: bytes) -> Tuple[str, str]:
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except Exception as exc:  # noqa: BLE001
        return "", f"OCR support unavailable: {exc}"

    images = convert_from_bytes(data)
    ocr_text = "\n".join(pytesseract.image_to_string(img) for img in images)
    return ocr_text, ""


def _extract_docx_text(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"--- Table {table_index} ---")
        for row in table.rows:
            row_cells = [cell.text for cell in row.cells]
            parts.append(_format_table([row_cells]))
    return "\n".join(parts)


def _read_uploaded_text(
    uploaded_file: Any,
    use_ocr: bool,
) -> Tuple[str, bytes, str]:
    if uploaded_file is None:
        return "", b"", ""
    data = uploaded_file.getvalue()
    name = (uploaded_file.name or "").lower()
    _, ext = os.path.splitext(name)
    if ext in {".txt", ".md", ".json", ".csv"}:
        return data.decode("utf-8", errors="replace"), data, ""
    if ext == ".pdf":
        try:
            text = _extract_pdf_text(data)
            if text.strip() or not use_ocr:
                return text, data, ""
            ocr_text, ocr_error = _ocr_pdf_text(data)
            if ocr_error:
                return "", data, ocr_error
            return ocr_text, data, ""
        except Exception as exc:  # noqa: BLE001
            return "", data, f"PDF support unavailable: {exc}"
    if ext == ".docx":
        try:
            return _extract_docx_text(data), data, ""
        except Exception as exc:  # noqa: BLE001
            return "", data, f"DOCX support unavailable: {exc}"
    return "", data, "Unsupported file type."


st.set_page_config(page_title="LangExtract Local App", layout="wide")
st.title("LangExtract Local App")
st.caption("Run LangExtract locally with your own model credentials or Ollama.")

with st.sidebar:
    st.header("Model Settings")
    provider = st.selectbox(
        "Provider",
        ["Ollama (local)", "Gemini", "OpenAI", "Other"],
        index=0,
    )
    default_model_ids = {
        "Ollama (local)": "gemma2:2b",
        "Gemini": "gemini-2.5-flash",
        "OpenAI": "gpt-4o",
        "Other": "",
    }
    model_id = st.text_input(
        "Model ID",
        value=default_model_ids[provider],
        key=f"model_id_{provider}",
    )
    if provider == "OpenAI":
        api_key_default = os.environ.get("OPENAI_API_KEY", "")
        api_key_help = "Defaults to OPENAI_API_KEY."
    else:
        api_key_default = os.environ.get("LANGEXTRACT_API_KEY", "")
        api_key_help = "Defaults to LANGEXTRACT_API_KEY."
    api_key = st.text_input(
        "API Key (optional)",
        value=api_key_default,
        type="password",
        help=api_key_help,
        disabled=provider == "Ollama (local)",
        key=f"api_key_{provider}",
    )
    if provider == "Ollama (local)":
        model_url = st.text_input("Ollama URL", value="http://localhost:11434")
    else:
        model_url = ""

    st.subheader("Advanced")
    if provider == "OpenAI":
        st.caption("OpenAI requires fence_output=true and schema constraints off.")
        fence_output = st.checkbox("Fence output", value=True, disabled=True)
        use_schema_constraints = st.checkbox(
            "Use schema constraints",
            value=False,
            disabled=True,
        )
    elif provider == "Ollama (local)":
        fence_output = st.checkbox("Fence output", value=False, disabled=True)
        use_schema_constraints = st.checkbox(
            "Use schema constraints",
            value=False,
            disabled=True,
        )
    else:
        fence_output = st.checkbox("Fence output", value=True)
        use_schema_constraints = st.checkbox("Use schema constraints", value=True)

    with st.expander("Diagnostics"):
        st.write(f"Python: {sys.executable}")
        st.write(f"Version: {sys.version.split()[0]}")
        if st.button("Check OCR modules"):
            try:
                import pdf2image  # noqa: F401
                st.success("pdf2image: OK")
            except Exception as exc:  # noqa: BLE001
                st.error(f"pdf2image: {exc}")
            try:
                import pytesseract  # noqa: F401
                st.success("pytesseract: OK")
            except Exception as exc:  # noqa: BLE001
                st.error(f"pytesseract: {exc}")

st.subheader("Extraction Input")
use_pdf_ocr = st.checkbox(
    "Use OCR for PDFs if no text is found",
    value=True,
)
uploaded_file = st.file_uploader(
    "Upload document (txt, md, json, csv, pdf, docx)",
    type=["txt", "md", "json", "csv", "pdf", "docx"],
)
uploaded_text, uploaded_bytes, upload_error = _read_uploaded_text(
    uploaded_file,
    use_ocr=use_pdf_ocr,
)
if upload_error:
    st.error(upload_error)

if "input_text" not in st.session_state:
    st.session_state["input_text"] = ""

reload_uploaded = st.button(
    "Reload from upload",
    disabled=uploaded_file is None,
)

if uploaded_file is not None:
    st.caption(
        f"Loaded {uploaded_file.name} ({len(uploaded_bytes)} bytes)."
    )
    uploaded_hash = hashlib.sha256(uploaded_bytes).hexdigest()
    st.session_state["uploaded_text"] = uploaded_text
    if reload_uploaded:
        if uploaded_text:
            st.session_state["input_text"] = uploaded_text
        else:
            st.warning("No text could be extracted from this file.")
    elif uploaded_text and st.session_state.get("uploaded_hash") != uploaded_hash:
        st.session_state["input_text"] = uploaded_text
        st.session_state["uploaded_hash"] = uploaded_hash

input_text = st.text_area(
    "Text to extract from",
    height=240,
    placeholder="Paste or type your document here...",
    key="input_text",
)
prompt_description = st.text_area(
    "Prompt description",
    height=120,
    placeholder="Describe what to extract (entities, relationships, etc.)...",
)

default_examples = json.dumps(
    [
        {
            "input": "John Smith paid $120 to ACME Corp on 2024-02-01.",
            "output": {
                "payer": "John Smith",
                "amount": "$120",
                "payee": "ACME Corp",
                "date": "2024-02-01",
            },
        }
    ],
    indent=2,
)
examples_raw = st.text_area(
    "Examples (JSON array)",
    value=default_examples,
    height=200,
)

col_left, col_right = st.columns([1, 1], gap="large")
with col_left:
    run = st.button("Run extraction", type="primary")

if run:
    if not input_text.strip() or not prompt_description.strip():
        st.error("Please provide both input text and a prompt description.")
    elif provider in {"Gemini", "OpenAI"} and not api_key:
        st.error("Provide an API key for the selected provider.")
        st.stop()
    else:
        try:
            examples = _deserialize_examples(examples_raw)
        except Exception as exc:  # noqa: BLE001
            st.error(f"Invalid examples JSON: {exc}")
            st.stop()

        language_model_params: dict[str, Any] = {}
        model_kwargs: dict[str, Any] = {
            "text_or_documents": input_text,
            "prompt_description": prompt_description,
            "examples": examples,
            "model_id": model_id,
            "fence_output": fence_output,
            "use_schema_constraints": use_schema_constraints,
        }

        if provider == "Ollama (local)":
            model_kwargs["model_url"] = model_url
            model_kwargs["fence_output"] = False
            model_kwargs["use_schema_constraints"] = False
        elif api_key:
            model_kwargs["api_key"] = api_key

        if language_model_params:
            model_kwargs["language_model_params"] = language_model_params

        with st.spinner("Running extraction..."):
            result = lx.extract(**model_kwargs)

        st.success("Extraction complete.")
        st.subheader("Result")
        jsonable = _result_to_jsonable(result)
        if isinstance(jsonable, (dict, list)):
            st.json(jsonable)
        else:
            st.write(jsonable)
